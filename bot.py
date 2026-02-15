import html
import logging
import os
import tempfile
import time
from io import BytesIO
from dataclasses import dataclass
from typing import Iterable, List, Optional, Tuple

from docx import Document
from dotenv import load_dotenv
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from telegram import InputFile, Update
from telegram.constants import ParseMode
from telegram.ext import ApplicationBuilder, ContextTypes, MessageHandler, filters

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 14
FIRST_LINE_INDENT_CM = 1.25
LINE_SPACING_PT = 18
DOCX_DEFAULT_FILENAME = "message.docx"
DOCX_FILENAME_MAX_DEFAULT = 60
ALLOWED_USER_IDS_ENV = "ALLOWED_USER_IDS"

TELEGRAM_MAX_LEN = 4096

SUPPORTED_ENTITY_TYPES = {
    "bold",
    "italic",
    "underline",
    "strikethrough",
    "code",
    "pre",
    "text_link",
    "url",
    "text_mention",
}


@dataclass(frozen=True)
class TextStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    code: bool = False
    link: Optional[str] = None


@dataclass(frozen=True)
class StyledSegment:
    text: str
    style: TextStyle


class HourlyHttpxFilter(logging.Filter):
    def __init__(self, interval_seconds: int = 3600) -> None:
        super().__init__()
        self.interval_seconds = interval_seconds
        self._last_emit = 0.0

    def filter(self, record: logging.LogRecord) -> bool:
        if record.name not in {"httpx", "httpcore"}:
            return True
        msg = record.getMessage()
        if "HTTP Request" not in msg:
            return True
        now = time.time()
        if now - self._last_emit >= self.interval_seconds:
            self._last_emit = now
            return True
        return False


def _sanitize_filename(value: str) -> str:
    cleaned = "".join(ch for ch in value if ch.isalnum() or ch in (" ", "-", "_"))
    cleaned = " ".join(cleaned.split())
    return cleaned.strip(" ._-")


def _derive_filename_from_text(text: str, max_len: int) -> str:
    if not text:
        return ""
    cutoff = len(text)
    newline_idx = text.find("\n")
    if newline_idx != -1:
        cutoff = min(cutoff, newline_idx)
    dot_idx = text.find(".")
    if dot_idx != -1:
        cutoff = min(cutoff, dot_idx)
    snippet = text[:cutoff].strip()
    if not snippet:
        return ""
    cleaned = _sanitize_filename(snippet)
    if not cleaned:
        return ""
    if max_len > 0:
        cleaned = cleaned[:max_len].rstrip(" ._-")
    return cleaned


def _parse_allowed_user_ids(value: Optional[str]) -> Optional[set[int]]:
    if not value:
        return None
    ids = set()
    for part in value.split(","):
        part = part.strip()
        if not part:
            continue
        try:
            ids.add(int(part))
        except ValueError:
            logging.warning("Invalid user id in %s: %s", ALLOWED_USER_IDS_ENV, part)
    return ids or None


def _is_user_allowed(user_id: Optional[int], allowed: Optional[set[int]]) -> bool:
    if allowed is None:
        return True
    if user_id is None:
        return False
    return user_id in allowed


def utf16_to_index(text: str, offset: int) -> int:
    if offset <= 0:
        return 0
    count = 0
    for i, ch in enumerate(text):
        if count == offset:
            return i
        count += 2 if ord(ch) > 0xFFFF else 1
        if count > offset:
            return i + 1
    return len(text)


def _is_on(el) -> bool:
    if el is None:
        return False
    val = el.get(qn("w:val"))
    if val is None:
        return True
    return val.lower() not in {"false", "0", "none"}


def _escape_attr(value: str) -> str:
    return html.escape(value, quote=True)


def _escape_html(value: str) -> str:
    return html.escape(value, quote=False)


def apply_paragraph_format(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(LINE_SPACING_PT)


def apply_run_format(run, style: TextStyle) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)

    run.bold = style.bold
    run.italic = style.italic
    run.underline = style.underline
    run.font.strike = style.strike


def add_hyperlink(paragraph, text: str, url: str, style: TextStyle) -> None:
    if not text:
        return
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.append(r_fonts)

    if style.bold:
        r_pr.append(OxmlElement("w:b"))
    if style.italic:
        r_pr.append(OxmlElement("w:i"))
    if style.underline:
        r_pr.append(OxmlElement("w:u"))
    if style.strike:
        r_pr.append(OxmlElement("w:strike"))

    size = str(int(FONT_SIZE_PT * 2))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), size)
    r_pr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), size)
    r_pr.append(sz_cs)

    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_segments_from_entities(text: str, entities) -> List[StyledSegment]:
    if not entities:
        return [StyledSegment(text, TextStyle())]

    normalized = []
    for ent in entities:
        if ent.type not in SUPPORTED_ENTITY_TYPES:
            continue
        start = utf16_to_index(text, ent.offset)
        end = utf16_to_index(text, ent.offset + ent.length)
        if start >= end:
            continue
        url = None
        if ent.type == "text_link":
            url = ent.url
        elif ent.type == "url":
            url = text[start:end]
        elif ent.type == "text_mention" and ent.user:
            url = f"tg://user?id={ent.user.id}"
        normalized.append({
            "type": ent.type,
            "start": start,
            "end": end,
            "url": url,
        })

    if not normalized:
        return [StyledSegment(text, TextStyle())]

    events = []
    for ent in normalized:
        events.append((ent["start"], 1, ent))
        events.append((ent["end"], 0, ent))
    events.sort(key=lambda x: (x[0], x[1]))  # end before start

    active = {
        "bold": 0,
        "italic": 0,
        "underline": 0,
        "strikethrough": 0,
        "code": 0,
        "link": [],
    }

    segments: List[StyledSegment] = []
    cursor = 0
    idx = 0
    while idx < len(events):
        offset = events[idx][0]
        if offset > cursor:
            seg_text = text[cursor:offset]
            style = TextStyle(
                bold=active["bold"] > 0,
                italic=active["italic"] > 0,
                underline=active["underline"] > 0,
                strike=active["strikethrough"] > 0,
                code=active["code"] > 0,
                link=active["link"][-1] if active["link"] else None,
            )
            segments.append(StyledSegment(seg_text, style))
            cursor = offset

        while idx < len(events) and events[idx][0] == offset:
            _, kind, ent = events[idx]
            delta = 1 if kind == 1 else -1
            etype = ent["type"]
            if etype == "pre":
                etype = "code"
            if etype in active:
                active[etype] += delta
            elif etype in {"text_link", "url", "text_mention"}:
                if delta > 0:
                    if ent["url"]:
                        active["link"].append(ent["url"])
                else:
                    if ent["url"] in active["link"]:
                        for i in range(len(active["link"]) - 1, -1, -1):
                            if active["link"][i] == ent["url"]:
                                active["link"].pop(i)
                                break
            idx += 1

    if cursor < len(text):
        style = TextStyle(
            bold=active["bold"] > 0,
            italic=active["italic"] > 0,
            underline=active["underline"] > 0,
            strike=active["strikethrough"] > 0,
            code=active["code"] > 0,
            link=active["link"][-1] if active["link"] else None,
        )
        segments.append(StyledSegment(text[cursor:], style))

    return segments


def telegram_text_to_docx(text: str, entities) -> Document:
    doc = Document()
    if not doc.paragraphs:
        para = doc.add_paragraph()
    else:
        para = doc.paragraphs[0]
    apply_paragraph_format(para)

    segments = build_segments_from_entities(text, entities)

    for seg in segments:
        parts = seg.text.split("\n")
        for i, part in enumerate(parts):
            if part:
                if seg.style.link:
                    add_hyperlink(para, part, seg.style.link, seg.style)
                else:
                    run = para.add_run(part)
                    apply_run_format(run, seg.style)
            if i < len(parts) - 1:
                para = doc.add_paragraph()
                apply_paragraph_format(para)

    return doc


def _run_text_from_xml(r) -> str:
    parts = []
    for child in r:
        if child.tag == qn("w:t"):
            parts.append(child.text or "")
        elif child.tag == qn("w:tab"):
            parts.append("    ")
        elif child.tag == qn("w:br"):
            parts.append("\n")
    return "".join(parts)


def _run_format_from_xml(r) -> TextStyle:
    r_pr = r.find(qn("w:rPr"))
    bold = italic = underline = strike = False
    if r_pr is not None:
        bold = _is_on(r_pr.find(qn("w:b")))
        italic = _is_on(r_pr.find(qn("w:i")))
        u_el = r_pr.find(qn("w:u"))
        if u_el is not None:
            underline = _is_on(u_el)
        strike = _is_on(r_pr.find(qn("w:strike"))) or _is_on(r_pr.find(qn("w:dstrike")))
    return TextStyle(bold=bold, italic=italic, underline=underline, strike=strike)


def iter_paragraph_runs_with_links(paragraph) -> Iterable[Tuple[str, TextStyle, Optional[str]]]:
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            text = _run_text_from_xml(child)
            if text:
                yield text, _run_format_from_xml(child), None
        elif child.tag == qn("w:hyperlink"):
            r_id = child.get(qn("r:id"))
            url = None
            if r_id:
                rel = paragraph.part.rels.get(r_id)
                if rel:
                    url = rel.target_ref
            for r in child.findall(qn("w:r")):
                text = _run_text_from_xml(r)
                if text:
                    yield text, _run_format_from_xml(r), url


def wrap_html(text: str, style: TextStyle, link: Optional[str]) -> str:
    if not text:
        return ""
    escaped = _escape_html(text)
    if link:
        escaped = f'<a href="{_escape_attr(link)}">{escaped}</a>'
    if style.bold:
        escaped = f"<b>{escaped}</b>"
    if style.italic:
        escaped = f"<i>{escaped}</i>"
    if style.underline:
        escaped = f"<u>{escaped}</u>"
    if style.strike:
        escaped = f"<s>{escaped}</s>"
    return escaped


def docx_to_telegram_html(docx_path: str) -> str:
    doc = Document(docx_path)
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph._p is None or len(paragraph._p) == 0:
            paragraphs.append("")
            continue
        parts = []
        for text, style, link in iter_paragraph_runs_with_links(paragraph):
            parts.append(wrap_html(text, style, link))
        paragraphs.append("".join(parts))
    return "\n".join(paragraphs)


def split_telegram_message(text: str, limit: int = TELEGRAM_MAX_LEN) -> List[str]:
    if len(text) <= limit:
        return [text]
    parts = []
    buffer = ""
    for para in text.split("\n"):
        candidate = para if not buffer else buffer + "\n" + para
        if len(candidate) > limit:
            if buffer:
                parts.append(buffer)
                buffer = para
            else:
                for i in range(0, len(para), limit):
                    parts.append(para[i : i + limit])
                buffer = ""
        else:
            buffer = candidate
    if buffer:
        parts.append(buffer)
    return parts


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    if not message or not message.text:
        return
    if not _is_user_allowed(
        message.from_user.id if message.from_user else None,
        context.application.bot_data.get("allowed_user_ids"),
    ):
        return

    user_id = message.from_user.id if message.from_user else None
    preview = message.text.replace("\n", " ")[:120]
    logging.info("IN text from %s: %s", user_id, preview)

    text = message.text
    entities = message.entities or []

    doc = telegram_text_to_docx(text, entities)

    max_len = int(os.getenv("DOCX_FILENAME_MAX", str(DOCX_FILENAME_MAX_DEFAULT)))
    derived = _derive_filename_from_text(text, max_len)
    filename = derived or os.getenv("DOCX_FILENAME", DOCX_DEFAULT_FILENAME)
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    await message.reply_document(document=InputFile(buffer, filename=filename))
    logging.info("OUT docx to %s: %s", user_id, filename)


def _media_types(message) -> List[str]:
    types = []
    if message.photo:
        types.append("photo")
    if message.video:
        types.append("video")
    if message.audio:
        types.append("audio")
    if message.voice:
        types.append("voice")
    if message.video_note:
        types.append("video_note")
    if message.animation:
        types.append("animation")
    if message.sticker:
        types.append("sticker")
    if message.contact:
        types.append("contact")
    if message.location:
        types.append("location")
    if message.venue:
        types.append("venue")
    if message.document:
        types.append("document")
    return types


async def handle_caption(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    if not message or not message.caption:
        return
    if not _is_user_allowed(
        message.from_user.id if message.from_user else None,
        context.application.bot_data.get("allowed_user_ids"),
    ):
        return

    user_id = message.from_user.id if message.from_user else None
    media = _media_types(message)
    if media:
        logging.warning("IN media from %s ignored: %s", user_id, ",".join(media))

    preview = message.caption.replace("\n", " ")[:120]
    logging.info("IN caption from %s: %s", user_id, preview)

    text = message.caption
    entities = message.caption_entities or []
    doc = telegram_text_to_docx(text, entities)

    max_len = int(os.getenv("DOCX_FILENAME_MAX", str(DOCX_FILENAME_MAX_DEFAULT)))
    derived = _derive_filename_from_text(text, max_len)
    filename = derived or os.getenv("DOCX_FILENAME", DOCX_DEFAULT_FILENAME)
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    await message.reply_document(document=InputFile(buffer, filename=filename))
    logging.info("OUT docx to %s: %s", user_id, filename)


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    if not message or not message.document:
        return
    if not _is_user_allowed(
        message.from_user.id if message.from_user else None,
        context.application.bot_data.get("allowed_user_ids"),
    ):
        return

    user_id = message.from_user.id if message.from_user else None
    doc = message.document
    filename = doc.file_name or "document"
    ext = os.path.splitext(filename)[1].lower()
    logging.info("IN doc from %s: %s", user_id, filename)

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, filename)
        file = await doc.get_file()
        await file.download_to_drive(input_path)

        try:
            if ext != ".docx":
                await message.reply_text("Пожалуйста, отправьте документ .docx")
                return

            html_text = docx_to_telegram_html(input_path)
            if not html_text:
                await message.reply_text("Документ пустой или не содержит читаемого текста.")
                return

            for chunk in split_telegram_message(html_text):
                await message.reply_text(chunk, parse_mode=ParseMode.HTML, disable_web_page_preview=True)
            logging.info("OUT text to %s: %d chars", user_id, len(html_text))
        except Exception as exc:
            logging.exception("Failed to process document")
            await message.reply_text(f"Ошибка обработки документа: {exc}")


async def handle_unsupported_media(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    message = update.effective_message
    if not message:
        return
    if not _is_user_allowed(
        message.from_user.id if message.from_user else None,
        context.application.bot_data.get("allowed_user_ids"),
    ):
        return
    if message.caption:
        return
    media = _media_types(message)
    if media:
        user_id = message.from_user.id if message.from_user else None
        logging.error("Unsupported media from %s (no caption): %s", user_id, ",".join(media))


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    httpx_logger = logging.getLogger("httpx")
    httpx_logger.addFilter(HourlyHttpxFilter())
    httpcore_logger = logging.getLogger("httpcore")
    httpcore_logger.addFilter(HourlyHttpxFilter())

    load_dotenv()
    allowed_user_ids = _parse_allowed_user_ids(os.getenv(ALLOWED_USER_IDS_ENV))
    token = os.getenv("BOT_TOKEN")
    if not token:
        raise RuntimeError("BOT_TOKEN environment variable is required")

    application = ApplicationBuilder().token(token).build()
    application.bot_data["allowed_user_ids"] = allowed_user_ids

    application.add_handler(MessageHandler(filters.CAPTION & ~filters.COMMAND, handle_caption))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_handler(
        MessageHandler(
            (filters.PHOTO | filters.VIDEO | filters.AUDIO | filters.VOICE | filters.ANIMATION | filters.VIDEO_NOTE | filters.STICKER),
            handle_unsupported_media,
        )
    )

    logging.info("Bot started")
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
